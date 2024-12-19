import streamlit as st
from docx import Document
import re
import requests
import os
from dotenv import load_dotenv
import msal
from openpyxl import load_workbook
from datetime import datetime
import calendar

# Set page configuration with a favicon
st.set_page_config(
    page_title="Prevista Invoice Automation",
    page_icon="https://lirp.cdn-website.com/d8120025/dms3rep/multi/opt/social-image-88w.png", 
    layout="centered"  # "centered" or "wide"
)

# Load environment variables from .env file
load_dotenv()

# Fetch credentials from environment variables
CLIENT_ID = os.getenv("CLIENT_ID")
CLIENT_SECRET = os.getenv("CLIENT_SECRET")
TENANT_ID = os.getenv("TENANT_ID")
DRIVE_ID = os.getenv("DRIVE_ID")

# Authenticate and acquire an access token
def acquire_access_token():
    app = msal.ConfidentialClientApplication(
        client_id=CLIENT_ID,
        client_credential=CLIENT_SECRET,
        authority=f"https://login.microsoftonline.com/{TENANT_ID}",
    )
    result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
    if "access_token" in result:
        return result["access_token"]
    else:
        print("Failed to acquire token")
        print(result.get("error"))
        print(result.get("error_description"))
        exit()

ACCESS_TOKEN = acquire_access_token() 

##############################
# Functions
##############################
    
def get_or_create_month_folder(access_token, drive_id, parent_folder_path):
    """
    Check if the current month's folder exists in the parent folder.
    If it doesn't exist, create the folder. Return the folder name as a string.

    Args:
        access_token (str): OAuth2 access token for authentication.
        drive_id (str): The ID of the drive to query.
        parent_folder_path (str): The path to the parent folder containing month folders.

    Returns:
        str: The name of the current month's folder.
    """
    # Get the current month and year
    current_month_name = calendar.month_name[datetime.now().month]
    current_year_suffix = str(datetime.now().year)[-2:]  # Last 2 digits of the year
    next_folder_index = datetime.now().month-6  # Assuming folders follow a numeric sequence (e.g., "5. November 24")
    current_month_folder_name = f"{next_folder_index}. {current_month_name} {current_year_suffix}"

    headers = {"Authorization": f"Bearer {access_token}"}
    parent_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{parent_folder_path}:/children"
    
    # Fetch items in the parent folder
    response = requests.get(parent_url, headers=headers)

    if response.status_code == 200:
        items = response.json().get("value", [])
        # Check if the current month's folder exists
        for item in items:
            if "folder" in item and current_month_folder_name in item["name"]:
                return current_month_folder_name  # Return folder name if it exists

        # Folder does not exist, create it
        create_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{parent_folder_path}:/children"
        folder_data = {
            "name": current_month_folder_name,
            "folder": {},
            "@microsoft.graph.conflictBehavior": "rename"
        }
        create_response = requests.post(create_url, headers=headers, json=folder_data)
        if create_response.status_code == 201:
            print(f"Folder '{current_month_folder_name}' created successfully.")
            return current_month_folder_name  # Return folder name after creation
        else:
            print(f"Error creating folder '{current_month_folder_name}': {create_response.status_code}")
            create_response.raise_for_status()
    else:
        print(f"Error fetching parent folder: {response.status_code}")
        response.raise_for_status()


def process_employee_folder(access_token, drive_id, parent_folder_path, employee_name, invoice_file_path, optional_files=[]):
    """
    Check if the folder exists for the employee. If not, create it.
    Move the uploaded invoice and optional files to the folder.

    Args:
        access_token (str): OAuth2 access token for authentication.
        drive_id (str): The ID of the drive to query.
        parent_folder_path (str): The path to the parent folder containing employee folders.
        employee_name (str): The employee name extracted from the invoice.
        invoice_file_path (str): Path to the uploaded invoice file.
        optional_files (list): List of optional files to upload.

    Returns:
        str: Status message indicating success or failure.
    """
    logs=[]
    # Fix name case (Title Case)
    formatted_name = employee_name.title()  # E.g., "muhammad osama" -> "Muhammad Osama"

    headers = {"Authorization": f"Bearer {access_token}"}
    parent_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{parent_folder_path}:/children"

    # Fetch folders in the parent path
    response = requests.get(parent_url, headers=headers)

    if response.status_code == 200:
        items = response.json().get("value", [])
        # Check if the employee's folder exists
        for item in items:
            if "folder" in item and formatted_name == item["name"]:
                employee_folder_id = item["id"]
                break
        else:
            # Folder does not exist, create it
            create_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{parent_folder_path}:/children"
            folder_data = {
                "name": formatted_name,
                "folder": {},
                "@microsoft.graph.conflictBehavior": "rename"
            }
            create_response = requests.post(create_url, headers=headers, json=folder_data)

            if create_response.status_code == 201:
                employee_folder_id = create_response.json()["id"]
            else:
                return f"Error creating folder: {create_response.status_code}"

        # Upload the mandatory invoice
        upload_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{employee_folder_id}:/{invoice_file_path}:/content"
        with open(invoice_file_path, "rb") as f:
            upload_response = requests.put(upload_url, headers=headers, data=f)
        if upload_response.status_code != 201:
            if upload_response.status_code == 200:
                logs.append(f"Invoice {invoice_file_path} Already exist with the same name.")
            else:
                logs.append(f"Error uploading invoice: {upload_response.status_code}")

        # Upload optional files (if provided)
        if optional_files:
            for optional_file in optional_files:
                upload_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{employee_folder_id}:/{optional_file.name}:/content"
                with open(optional_file.name, "rb") as f:
                    upload_response = requests.put(upload_url, headers=headers, data=f)
                if upload_response.status_code != 201:
                    if upload_response.status_code == 200:
                        logs.append(f"Receipt {optional_file.name} Already exist with the same name.")
                    else:
                        logs.append(f"Error uploading file '{optional_file.name}': {upload_response.status_code}")

        logs.append(f"Files uploaded successfully to folder: {formatted_name}")
    else:
        logs.append(f"Error fetching parent folder: {response.status_code}")
    
    return logs


def find_master_sheet_path(access_token, drive_id, folder_path):
    """
    Find the master sheet (an .xlsx file with 'Invoices' in the name) in the specified folder
    and return its SharePoint file path.

    Args:
        access_token (str): OAuth2 access token for authentication.
        drive_id (str): The ID of the drive to query.
        folder_path (str): The path to the folder in the drive.

    Returns:
        str: The SharePoint file path of the master sheet.
    """
    # Get the list of files and folders
    list_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{folder_path}:/children"
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(list_url, headers=headers)

    if response.status_code == 200:
        files_and_folders = response.json().get("value", [])
    else:
        raise Exception(f"Error fetching files from folder: {response.status_code} - {response.text}")

    # Find the master sheet
    for item in files_and_folders:
        if item["name"].endswith(".xlsx") and "Invoices" in item["name"]:
            # Return the SharePoint file path
            return f"{folder_path}/{item['name']}"

    # If no file is found, raise an exception
    raise FileNotFoundError("No master sheet found (file with '.xlsx' and 'Invoices' in the name).")


def update_mastersheet_sharepoint(access_token, drive_id, file_path, employee_name, total, month="Jan-24"):
    """
    Update the master sheet in SharePoint by modifying only the required cell.

    Args:
        access_token (str): OAuth2 access token for authentication.
        drive_id (str): The ID of the SharePoint drive to query.
        file_path (str): The relative path to the file in SharePoint.
        employee_name (str): The name of the employee.
        total (float): The total to update for the given month.
        month (str): The month to update in the format 'Dec-24'.

    Returns:
        str: Status message indicating success or failure.
    """
    try:
        # Step 1: Download the file from SharePoint
        download_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{file_path}:/content"
        headers = {"Authorization": f"Bearer {access_token}"}
        response = requests.get(download_url, headers=headers)

        if response.status_code != 200:
            return f"Error downloading file: {response.status_code} - {response.text}"

        # Save the file locally
        local_file_path = "temp_master_sheet.xlsx"
        with open(local_file_path, "wb") as f:
            f.write(response.content)

        # Step 2: Update the file locally using the existing logic
        workbook = load_workbook(local_file_path)
        sheet = workbook.active  # Assuming the master sheet is the active sheet

        # Define the relevant row/column ranges
        start_row = 38
        end_row = 87
        name_column = "C"  # Column C (e.g., employee names)
        text_column = "B"  # Column B (e.g., 'STARFLEET  / Catalyst')
        month_headers_row = 7  # Row 7 contains month headers
        start_month_column = 8  # Column H (1-based index)

        # Convert the month input to a datetime object for comparison
        target_month = datetime.strptime(month, "%b-%y")

        # Match the month column
        for col in range(start_month_column, start_month_column + 12):  # Columns H to S
            cell_value = sheet.cell(row=month_headers_row, column=col).value
            if isinstance(cell_value, datetime):
                formatted_header = cell_value.strftime("%b-%y")
            else:
                formatted_header = cell_value

            if formatted_header == month:
                current_month_col = col
                break
        else:
            os.remove(local_file_path)
            return f"Error: Month '{month}' not found in master sheet."

        # Locate the employee's row
        for row in range(start_row, end_row + 1):
            if (
                sheet[f"{text_column}{row}"].value == "STARFLEET  / Catalyst" and
                sheet[f"{name_column}{row}"].value.strip() == employee_name.strip()
            ):
                # Update the cell for the current month
                sheet.cell(row=row, column=current_month_col).value = total
                break
        else:
            os.remove(local_file_path)
            return f"Error: Employee '{employee_name}' not found in the master sheet."

        # Save the updated workbook locally
        workbook.save(local_file_path)

        # Step 3: Upload the updated file back to SharePoint
        upload_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{file_path}:/content"
        with open(local_file_path, "rb") as f:
            upload_response = requests.put(upload_url, headers=headers, data=f)

        # Remove the local temporary file
        os.remove(local_file_path)

        if upload_response.status_code == 200:
            return f"Successfully updated total for '{employee_name}' in '{month}'."
        else:
            return f"Error uploading file back to SharePoint: {upload_response.status_code} - {upload_response.text}"

    except Exception as e:
        return f"An error occurred: {str(e)}"

def current_academic_year():
    # Determine the current academic year
    current_date = datetime.now()
    current_year = current_date.year
    current_month = current_date.month

    # Calculate the academic year
    if current_month >= 8:  # August to December
        start_year = current_year
        end_year = current_year + 1
    else:  # January to July
        start_year = current_year - 1
        end_year = current_year

    academic_year = f"{start_year}-{str(end_year)[-2:]}"  # E.g., "2024-25"

    return academic_year

def get_or_create_base_folder_path(access_token, drive_id):
    """
    Dynamically determine the base folder path for the academic year and create the folder
    in SharePoint if it doesn't exist.

    Args:
        access_token (str): OAuth2 access token for Microsoft Graph API.
        drive_id (str): The ID of the SharePoint drive.

    Returns:
        str: The SharePoint folder path.
    """
    academic_year = current_academic_year()
    base_folder_path = f"AEB Financial/{academic_year}/Invoices"

    # Check if the folder exists
    check_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{base_folder_path}"
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(check_url, headers=headers)

    if response.status_code == 404:  # Folder does not exist
        # Create the folder
        create_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{base_folder_path}:/children"
        folder_data = {
            "name": academic_year,
            "folder": {},
            "@microsoft.graph.conflictBehavior": "rename"
        }
        create_response = requests.post(create_url, headers=headers, json=folder_data)
        if create_response.status_code == 201:
            print(f"Folder '{base_folder_path}' created successfully.")
        else:
            raise Exception(f"Error creating folder: {create_response.status_code} - {create_response.text}")
    elif response.status_code != 200:
        raise Exception(f"Error checking folder: {response.status_code} - {response.text}")

    # Return the base folder path
    return base_folder_path


def list_files_by_path(access_token, drive_id, folder_path):
    """
    Lists files and folders in a specified path in OneDrive.

    Args:
        access_token (str): OAuth2 access token for authentication.
        drive_id (str): The ID of the drive to query.
        folder_path (str): The path to the folder in the drive.

    Returns:
        list: A list of dictionaries containing file/folder details (name, type, ID).
    """
    headers = {"Authorization": f"Bearer {access_token}"}
    url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{folder_path}:/children"
    response = requests.get(url, headers=headers)
    
    if response.status_code == 200:
        items = response.json().get("value", [])
        result = []
        for item in items:
            result.append({
                "name": item["name"],
                "type": "folder" if "folder" in item else "file",
                "id": item["id"]
            })
        return result
    else:
        print(f"Error fetching files by path: {response.status_code}")
        print(response.json())
        return []


def download_invoice_template():
    with open("resources/invoice_template.docx", "rb") as file:
        st.download_button(
            label="Download invoice template",
            data=file,
            file_name="invoice_template.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

def extract_total_from_invoice(docx_file):
    """
    Extract total amount from a DOCX invoice file
    """
    doc = Document(docx_file)
    total = 0
    
    # Search for total amount in document paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.lower()
        if 'total' in text:
            # Extract numbers from text
            numbers = re.findall(r'\d+\.?\d*', text)
            if numbers:
                # Take the last number found as total
                total = float(numbers[-1])
    return total

def extract_name_from_invoice(docx_file):
    """
    Extract the name from the 'Your Name' field in the DOCX invoice file
    """
    doc = Document(docx_file)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("YOUR NAME"):
            # Extract the name by splitting the text on the first tab or space
            parts = text.split("\t")
            if len(parts) > 1:
                return parts[1].strip()  # Return the name after the tab
            else:
                # If tabs are not used, fallback to space-based split
                parts = text.split()
                if len(parts) > 2:
                    return " ".join(parts[2:]).strip()  # Return everything after "YOUR NAME"
    return None  # Return None if "Your Name" field is not found


##############################
# Main
##############################

def main():
    # Initialize session state variables
    if "invoice_uploaded" not in st.session_state:
        st.session_state["invoice_uploaded"] = False
    if "processed" not in st.session_state:
        st.session_state["processed"] = False
    if "submitting" not in st.session_state:
        st.session_state["submitting"] = False

    # Logo space
    st.markdown(
        """
        <style>
            .logo {
                display: flex;
                justify-content: center;
                align-items: center;
                padding: 20px;
                border-bottom: 2px solid #f0f0f0;
                margin-bottom: 20px;
            }
        </style>
        """, unsafe_allow_html=True
    )

    st.image("resources/logo_removed_bg - enlarged.png", use_column_width=True)

    # Page title
    st.markdown(
        """
        <h2 style="text-align:center; color:#4CAF50;">Invoice Submission System</h2>
        """, unsafe_allow_html=True
    )

    # Download Invoice Template
    st.write("")
    st.divider()

    st.markdown(
        """
        <div style="text-align:left;">
            <p>Need an invoice template? Click below to download:</p>
        </div>
        """, unsafe_allow_html=True
    )
    download_invoice_template()
    st.divider()

    # File Upload Section
    st.markdown("<h3>Upload your Files</h3>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        invoice_file = st.file_uploader(
            "Upload Invoice (DOCX file)", 
            type=["docx"], 
            help="Upload your invoice here (mandatory)"
        )
    with col2:
        receipt_files = st.file_uploader(
            "Upload Expense Receipts (Optional, multiple)", 
            type=["jpg", "png", "pdf"], 
            accept_multiple_files=True,
            help="Upload your receipts here (optional)"
        )

    # Update state if invoice file is uploaded
    if invoice_file:
        st.session_state["invoice_uploaded"] = True
        st.success(f"Uploaded Invoice: {invoice_file.name}")
    else:
        st.session_state["invoice_uploaded"] = False

    # Display Uploaded Files
    if receipt_files:
        for receipt in receipt_files:
            st.success(f"Uploaded Receipt: {receipt.name}")

    # Process Button and Total Display
    st.markdown("<h3>Process & Submit Your Invoice</h3>", unsafe_allow_html=True)
    if st.button("[MUST] Please Click here to check if your name & total is correct"):
        if not st.session_state["invoice_uploaded"]:
            st.error("Please upload an invoice before processing!")
        else:
            st.session_state["processed"] = True  
            with st.spinner("Processing..."):
                # time.sleep(2)  # Simulating processing time
                st.session_state["name"] = extract_name_from_invoice(invoice_file)
                st.session_state["total"] = extract_total_from_invoice(invoice_file)
                st.success("Invoice processed successfully!")

                # Display Total (if processed)
                st.text(f"Name: {st.session_state.get('name', 0)}")
                st.text(f"Total: {st.session_state.get('total', 0)}")
                st.write("If the above name OR total is not correct, please contact us via email.")
                st.write("If the above name & total is correct, click Submit to complete your invoice submission:")

    # Submit Button Logic
    if st.session_state["invoice_uploaded"] and st.session_state["processed"]:
        if st.session_state["submitting"]:
            with st.spinner("Submitting..."):
                # time.sleep(5)  # Simulating submission

                ##############################
                # All Logic here from sharepoint file upload to update master sheet & Folder creation if doesn't exist
                ##############################

                # BASE_FOLDER_PATH = "AEB Financial/2024-25/Invoices"
                # Get or create the base folder path
                try:
                    BASE_FOLDER_PATH = get_or_create_base_folder_path(ACCESS_TOKEN, DRIVE_ID)
                    print(f"Base folder path: {BASE_FOLDER_PATH}")
                except Exception as e:
                    print(f"An error occurred: {e}")

                month_folder = get_or_create_month_folder(ACCESS_TOKEN, DRIVE_ID, BASE_FOLDER_PATH)
                FOLDER_PATH = f"{BASE_FOLDER_PATH}/{month_folder}/Catalyst"

                # Process employee folder
                ##########################
                process_employee_folder_result_message = process_employee_folder(
                    ACCESS_TOKEN,
                    DRIVE_ID,
                    FOLDER_PATH,
                    st.session_state["name"],
                    f"{invoice_file.name}",  # Mandatory invoice
                    optional_files=receipt_files
                )
                # Log for eamil!
                st.text("Log: Process employee folder")
                for message in process_employee_folder_result_message:
                    st.text("Log: " + f"{message}")

                # Process master sheet
                ######################
                master_FILE_PATH = find_master_sheet_path(ACCESS_TOKEN, DRIVE_ID, f"AEB Financial/{current_academic_year()}")
                EMPLOYEE_NAME = st.session_state["name"]
                TOTAL = st.session_state["total"]
                MONTH =  datetime.now().strftime("%b-%y")

                update_mastersheet_message = update_mastersheet_sharepoint(ACCESS_TOKEN, DRIVE_ID, master_FILE_PATH, EMPLOYEE_NAME, TOTAL, MONTH)
                # Log for eamil!
                st.text("Log: Process master sheet")
                st.text("Log: "+f"{update_mastersheet_message}")

                ##############################
                # ####################
                ##############################

                st.session_state["submitting"] = False
                st.success("Invoice Submitted Successfully!")
                

        else:
            if st.button("Submit"):
                st.session_state["submitting"] = True
                st.experimental_rerun()  # Rerun immediately to start the spinner
    else:
        st.button("Submit", disabled=True)


if __name__ == "__main__":
    main()

# Dev : https://linkedin.com/in/osamatech786