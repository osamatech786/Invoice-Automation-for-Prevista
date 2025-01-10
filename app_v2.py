import streamlit as st
from docx import Document
from io import BytesIO
from datetime import datetime, date, timedelta
import re
import requests
import os
from dotenv import load_dotenv
import msal
from openpyxl import load_workbook
import calendar
import pandas as pd
from docx.shared import Pt
from pytz import timezone, all_timezones
import shutil
from email.message import EmailMessage
import smtplib
import zipfile
import io

# Set page configuration with a favicon
st.set_page_config(
    page_title="Prevista Invoice Generator",
    page_icon="https://lirp.cdn-website.com/d8120025/dms3rep/multi/opt/social-image-88w.png", 
    layout="centered"  # "centered" or "wide"
)

# Load environment variables from .env file
load_dotenv()

# Fetch credentials from environment variables
CLIENT_ID = os.getenv("CLIENT_ID")
CLIENT_SECRET = os.getenv("CLIENT_SECRET")
TENANT_ID = os.getenv("TENANT_ID")
DRIVE_ID = os.getenv("DRIVE_ID")


# ========================
# Functions
# ========================

def fetch_calendar_events(access_token, employee_email):
    """
    Fetch calendar events for the current month for a given user.

    Parameters:
        access_token (str): Microsoft Graph API access token.
        employee_email (str): Email address of the user whose calendar to fetch.

    Returns:
        list: List of calendar events with start and end times within the current month.
    """
    # Calculate the start and end of the current month
    today = datetime.today()
    start_of_month = today.replace(day=1).isoformat()  # First day of the month
    last_day = calendar.monthrange(today.year, today.month)[1]  # Last day of the month
    end_of_month = today.replace(day=last_day, hour=23, minute=59, second=59).isoformat()  # End of the month

    # Set up headers with the access token
    headers = {"Authorization": f"Bearer {access_token}"}

    # Microsoft Graph API endpoint to fetch events
    url = f"https://graph.microsoft.com/v1.0/users/{employee_email}/calendar/events"
    params = {
        "startDateTime": start_of_month,
        "endDateTime": end_of_month,
        "$select": "id,start,end,subject,location"
    }

    # Make the API request
    response = requests.get(url, headers=headers, params=params)

    # Check response and return results
    if response.status_code == 200:
        events = response.json().get('value', [])
        processed_events = []

        for event in events:
            event_start = event['start']['dateTime']
            event_end = event['end']['dateTime']

            # Filter events strictly within the current month
            event_start_dt = datetime.fromisoformat(event_start)
            if event_start_dt.month == today.month and event_start_dt.year == today.year:
                processed_events.append({
                    "id": event['id'],
                    "title": event.get('subject', 'No Title'),
                    "start": event_start,
                    "end": event_end,
                    "location": event.get('location', {}).get('displayName', 'No Location')
                })
        
        return processed_events
    else:
        print(f"Error fetching events: {response.status_code} - {response.text}")
        return []


def validate_sessions(user_sessions, api_events, user_timezone_str):
    results = []
    user_timezone = timezone(user_timezone_str)  # Get the user's timezone

    for session in user_sessions:
        # Parse user-entered session time and localize it to user's timezone
        user_time_naive = datetime.strptime(f"{session['date']} {session['time']}", "%d-%m-%Y %H:%M:%S")
        user_time = user_timezone.localize(user_time_naive)  # Convert to timezone-aware datetime

        match_found = False
        for event in api_events:
            # Parse API event times and convert them to the user's timezone
            event_start = datetime.fromisoformat(event["start"].replace("Z", "+00:00")).astimezone(user_timezone)
            event_end = datetime.fromisoformat(event["end"].replace("Z", "+00:00")).astimezone(user_timezone)

            # Compare the localized times
            if event_start <= user_time <= event_end:
                match_found = True
                results.append({
                    "Session": session,
                    "Status": "Matched",
                    "Event": {
                        "title": event["title"],
                        "start": event_start.strftime("%d-%m-%Y %H:%M:%S"),
                        "end": event_end.strftime("%d-%m-%Y %H:%M:%S"),
                        "timezone": user_timezone_str
                    }
                })
                break

        if not match_found:
            results.append({"Session": session, "Status": "No Match Found", "Event": None})

    return results


# Authenticate and acquire an access token
def acquire_access_token():
    app = msal.ConfidentialClientApplication(
        client_id=CLIENT_ID,
        client_credential=CLIENT_SECRET,
        authority=f"https://login.microsoftonline.com/{TENANT_ID}",
    )
    result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
    if "access_token" in result:
        return result["access_token"]
    else:
        print("Failed to acquire token")
        print(result.get("error"))
        print(result.get("error_description"))
        exit()


# Define a function to calculate progress and percentage
def get_progress(step, total_steps=14):
    return int((step / total_steps) * 100)


# Function to replace placeholders in the template
def replace_placeholders(template_path, data, table_data):
    doc = Document(template_path)
    
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in data.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))  # Convert value to string
    
    # Replace table data
    for table in doc.tables:
        if "Date" in table.cell(0, 0).text:  # Find the target table by its header
            for i, row_data in enumerate(table_data):
                row_cells = table.add_row().cells
                row_cells[0].text = str(row_data["date"])
                row_cells[1].text = str(row_data["time_hours"])
                row_cells[2].text = str(row_data["activity"])
                row_cells[3].text = f"£{float(row_data['amount']):.2f}"  # Format amount as currency
                # Optionally format the text
                for cell in row_cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
            break
    return doc


def get_or_create_month_folder(access_token, drive_id, parent_folder_path):
    """
    Check if the current month's folder exists in the parent folder.
    If it doesn't exist, create the folder. Return the folder name as a string.

    Args:
        access_token (str): OAuth2 access token for authentication.
        drive_id (str): The ID of the drive to query.
        parent_folder_path (str): The path to the parent folder containing month folders.

    Returns:
        str: The name of the current month's folder.
    """
    # Get the current month and year
    current_month_name = calendar.month_name[datetime.now().month]
    current_year_suffix = str(datetime.now().year)[-2:]  # Last 2 digits of the year

    # Calculate the base index, starting from 1 for July
    base_month = 7  # July is the starting month
    current_month = datetime.now().month
    current_year = datetime.now().year
    start_year = 2024  # The year the sequence started

    # Calculate the index based on the starting point
    months_since_start = (current_year - start_year) * 12 + (current_month - base_month)
    next_folder_index = months_since_start + 1  # +1 to start the index from 1 for July
    
    current_month_folder_name = f"{next_folder_index}. {current_month_name} {current_year_suffix}"

    headers = {"Authorization": f"Bearer {access_token}"}
    parent_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{parent_folder_path}:/children"
    
    # Fetch items in the parent folder
    response = requests.get(parent_url, headers=headers)

    if response.status_code == 200:
        items = response.json().get("value", [])
        # Check if the current month's folder exists
        for item in items:
            if "folder" in item and current_month_folder_name in item["name"]:
                return current_month_folder_name  # Return folder name if it exists

        # Folder does not exist, create it
        create_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{parent_folder_path}:/children"
        folder_data = {
            "name": current_month_folder_name,
            "folder": {},
            "@microsoft.graph.conflictBehavior": "rename"
        }
        create_response = requests.post(create_url, headers=headers, json=folder_data)
        
        # for Catalyst folder
        # {month_folder}/Catalyst
        if create_response.status_code == 201:  # Successfully created the main folder
            main_folder_id = create_response.json()["id"]  # Get the ID of the main folder
            # Create the "Catalyst" folder inside the newly created main folder
            catalyst_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{main_folder_id}/children"
            catalyst_folder_data = {
                "name": "Catalyst",
                "folder": {},
                "@microsoft.graph.conflictBehavior": "rename"
            }
            catalyst_response = requests.post(catalyst_url, headers=headers, json=catalyst_folder_data)

            if catalyst_response.status_code == 201:
                print("Catalyst folder created successfully inside the main folder.")
            else:
                print(f"Error creating Catalyst folder: {catalyst_response.status_code}")
        
        if create_response.status_code == 201:
            print(f"Folder '{current_month_folder_name}' created successfully.")
            return current_month_folder_name  # Return folder name after creation
        else:
            print(f"Error creating folder '{current_month_folder_name}': {create_response.status_code}")
            create_response.raise_for_status()
    else:
        print(f"Error fetching parent folder in function 'get_or_create_month_folder': {response.status_code}")
        response.raise_for_status()


def process_employee_folder(access_token, drive_id, parent_folder_path, employee_name, invoice_file_path, optional_files=[]):
    """
    Check if the folder exists for the employee. If not, create it.
    Move the uploaded invoice and optional files to the folder.

    Args:
        access_token (str): OAuth2 access token for authentication.
        drive_id (str): The ID of the drive to query.
        parent_folder_path (str): The path to the parent folder containing employee folders.
        employee_name (str): The employee name extracted from the invoice.
        invoice_file_path (str): Path to the uploaded invoice file.
        optional_files (list): List of optional files to upload.

    Returns:
        str: Status message indicating success or failure.
    """
    logs=[]
    # Fix name case (Title Case)
    formatted_name = employee_name.title()  # E.g., "muhammad osama" -> "Muhammad Osama"

    headers = {"Authorization": f"Bearer {access_token}"}
    parent_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{parent_folder_path}:/children"

    # Fetch folders in the parent path
    response = requests.get(parent_url, headers=headers)

    if response.status_code == 200:
        items = response.json().get("value", [])
        # Check if the employee's folder exists
        for item in items:
            if "folder" in item and formatted_name == item["name"]:
                employee_folder_id = item["id"]
                break
        else:
            # Folder does not exist, create it
            create_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{parent_folder_path}:/children"
            folder_data = {
                "name": formatted_name,
                "folder": {},
                "@microsoft.graph.conflictBehavior": "rename"
            }
            create_response = requests.post(create_url, headers=headers, json=folder_data)

            if create_response.status_code == 201:
                employee_folder_id = create_response.json()["id"]
            else:
                return f"Error creating folder: {create_response.status_code}"

        # Upload the mandatory invoice
        upload_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{employee_folder_id}:/{invoice_file_path}:/content"
        with open(invoice_file_path, "rb") as f:
            upload_response = requests.put(upload_url, headers=headers, data=f)
        if upload_response.status_code != 201:
            if upload_response.status_code == 200:
                logs.append(f"Invoice {invoice_file_path} Already exist with the same name.")
            else:
                logs.append(f"Error uploading invoice: {upload_response.status_code}")

        # Upload optional files (if provided)
        if optional_files:
            for optional_file in optional_files:
                upload_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{employee_folder_id}:/{optional_file.name}:/content"
                with open(optional_file.name, "rb") as f:
                    upload_response = requests.put(upload_url, headers=headers, data=f)
                if upload_response.status_code != 201:
                    if upload_response.status_code == 200:
                        logs.append(f"Receipt {optional_file.name} Already exist with the same name.")
                    else:
                        logs.append(f"Error uploading file '{optional_file.name}': {upload_response.status_code}")

        logs.append(f"Files uploaded successfully to folder: {formatted_name}")
    else:
        logs.append(f"Error fetching parent folder in 'process_employee_folder': {response.status_code}")
    
    return logs


# Function to fetch and read recipients from the "Email" sheet of an Excel file
def fetch_recipients_from_sharepoint(access_token, drive_id):
    try:
        academic_year = current_academic_year()
        folder_path = f"AEB Financial/{academic_year}"
        file_path = find_master_sheet_path(access_token, drive_id, folder_path)

        # Download the Excel file
        download_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{file_path}:/content"
        headers = {"Authorization": f"Bearer {access_token}"}
        response = requests.get(download_url, headers=headers)

        if response.status_code != 200:
            raise Exception(f"Error downloading file: {response.status_code} - {response.text}")

        # Load the Excel content
        excel_data = pd.read_excel(BytesIO(response.content), sheet_name="Email")

        if "Email" not in excel_data.columns:
            raise ValueError("The 'Email' sheet must contain 'Email' column.")

        return list(zip(excel_data["Email"], excel_data["UTR"], excel_data["Name"], excel_data["Invoice Number"], excel_data["Centre Number"], excel_data["Pay Rate"], excel_data["Account Name"], excel_data["Branch Name"], excel_data["Sort Code"], excel_data["Account Number"], excel_data["JD"]))        
    
    except Exception as e:
        st.error(f"Error fetching recipients: {e}")
        return []

def find_master_sheet_path(access_token, drive_id, folder_path):
    """
    Find the master sheet (an .xlsx file with 'Invoices' in the name) in the specified folder
    and return its SharePoint file path.

    Args:
        access_token (str): OAuth2 access token for authentication.
        drive_id (str): The ID of the drive to query.
        folder_path (str): The path to the folder in the drive.

    Returns:
        str: The SharePoint file path of the master sheet.
    """
    # Get the list of files and folders
    list_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{folder_path}:/children"
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(list_url, headers=headers)

    if response.status_code == 200:
        files_and_folders = response.json().get("value", [])
    else:
        raise Exception(f"Error fetching files from folder: {response.status_code} - {response.text}")

    # Find the master sheet
    for item in files_and_folders:
        if item["name"].endswith(".xlsx") and "Invoices" in item["name"]:
            # Return the SharePoint file path
            return f"{folder_path}/{item['name']}"

    # If no file is found, raise an exception
    raise FileNotFoundError("No master sheet found (file with '.xlsx' and 'Invoices' in the name).")

def clean_name(name):
    """
    Remove bracketed text from a name.
    Args:
        name (str): The name with or without bracketed text.
    Returns:
        str: The cleaned name without any bracketed text.
    """
    import re
    return re.sub(r"\s*\(.*?\)", "", name).strip()


def update_mastersheet_sharepoint(access_token, drive_id, file_path, employee_name, total, month="Jan-25"):
    """
    Update the master sheet in SharePoint by modifying only the required cell using the Microsoft Graph API.
    Includes detailed debugging for header values and handles Excel serial dates.

    Args:
        access_token (str): OAuth2 access token for authentication.
        drive_id (str): The ID of the SharePoint drive to query.
        file_path (str): The relative path to the file in SharePoint.
        employee_name (str): The name of the employee.
        total (float): The total to update for the given month.
        month (str): The month to update in the format 'MMM-yy'.

    Returns:
        str: Status message indicating success or failure.
    """
    try:
        # Step 1: Fetch file ID
        print(f"Fetching file ID for path: {file_path}")
        get_file_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{file_path}"
        headers = {"Authorization": f"Bearer {access_token}"}
        response = requests.get(get_file_url, headers=headers)

        if response.status_code != 200:
            print(f"Error fetching file ID: {response.status_code} - {response.text}")
            return f"Error fetching file ID: {response.status_code} - {response.text}"

        file_id = response.json().get("id")
        print(f"File ID: {file_id}")

        # Step 2: Get the first visible worksheet
        print("Fetching list of worksheets...")
        sheets_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{file_id}/workbook/worksheets"
        sheets_response = requests.get(sheets_url, headers=headers)

        if sheets_response.status_code != 200:
            print(f"Error fetching worksheets: {sheets_response.status_code} - {sheets_response.text}")
            return f"Error fetching worksheets: {sheets_response.status_code} - {sheets_response.text}"

        sheets = sheets_response.json().get("value", [])
        visible_sheet_name = None
        for sheet in sheets:
            if sheet.get("visibility", "") == "Visible":
                visible_sheet_name = sheet["name"]
                print(f"First visible sheet found: {visible_sheet_name}")
                break

        if not visible_sheet_name:
            print("Error: No visible sheets found in the workbook.")
            return "Error: No visible sheets found in the workbook."

        # Step 3: Fetch month headers
        month_headers_row = 7
        month_headers_range = f"H{month_headers_row}:S{month_headers_row}"  # Adjusted range based on screenshot
        print(f"Fetching month headers from range: {month_headers_range}")

        headers_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{file_id}/workbook/worksheets('{visible_sheet_name}')/range(address='{month_headers_range}')"
        headers_response = requests.get(headers_url, headers=headers)

        if headers_response.status_code != 200:
            print(f"Error fetching month headers: {headers_response.status_code} - {headers_response.text}")
            return f"Error fetching month headers: {headers_response.status_code} - {headers_response.text}"

        headers_values = headers_response.json().get("values", [[]])[0]
        print("Raw month headers fetched:", headers_values)

        # Step 4: Match the provided month with headers (convert serial dates to datetime)
        current_month_col = None
        excel_start_date = datetime(1899, 12, 30)  # Excel starts from 30-Dec-1899
        for col_index, header in enumerate(headers_values, start=8):  # Adjust starting index if necessary
            print(f"Checking header: {header} at column {col_index}")
            try:
                # Convert Excel serial number to datetime
                header_date = excel_start_date + timedelta(days=int(header))
                formatted_header = header_date.strftime("%b-%y")
                if formatted_header == month:
                    current_month_col = col_index
                    print(f"Matched date header: {header_date} (formatted: {formatted_header}) at column {col_index}")
                    break
            except (ValueError, TypeError) as e:
                print(f"Skipping invalid header format: {header}, error: {e}")

        if current_month_col is None:
            print(f"Error: Month '{month}' not found in master sheet.")
            return f"Error: Month '{month}' not found in master sheet."

        print(f"Matched month column: {current_month_col}")

        # Step 5: Fetch employee data
        name_column_range = f"B1:C147"
        print(f"Fetching employee data from range: {name_column_range}")

        name_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{file_id}/workbook/worksheets('{visible_sheet_name}')/range(address='{name_column_range}')"
        name_response = requests.get(name_url, headers=headers)

        if name_response.status_code != 200:
            print(f"Error fetching employee rows: {name_response.status_code} - {name_response.text}")
            return f"Error fetching employee rows: {name_response.status_code} - {name_response.text}"

        name_values = name_response.json().get("values", [])
        print("Employee data fetched:", name_values)

        target_row = None
        for row_index, row in enumerate(name_values, start=1):
            print(f"Checking row {row_index}: {row}")
            # if len(row) > 1 and row[1].strip().lower() == employee_name.strip().lower():
            if len(row) > 1 and clean_name(row[1].strip().lower()) == clean_name(employee_name.strip().lower()):
                st.write(clean_name(row[1].strip().lower()))
                st.write(row[1].strip().lower())
                st.write(clean_name(employee_name.strip().lower()))
                st.write(employee_name.strip().lower())
                target_row = row_index
                print(f"Matched employee: {employee_name} at row {row_index}")
                break

        if target_row is None:
            print(f"Error: Employee '{employee_name}' not found in the master sheet.")
            return f"Error: Employee '{employee_name}' not found in the master sheet."

        print(f"Matched employee row: {target_row}")

        # Step 6: Update the target cell
        target_cell_address = f"{chr(64 + current_month_col)}{target_row}"
        print(f"Updating cell: {target_cell_address} with value: {total}")

        update_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{file_id}/workbook/worksheets('{visible_sheet_name}')/range(address='{target_cell_address}')"
        update_data = {"values": [[total]]}

        update_response = requests.patch(update_url, headers=headers, json=update_data)

        if update_response.status_code == 200:
            print(f"Successfully updated total for '{employee_name}' in '{month}'.")
            return f"Successfully updated total for '{employee_name}' in '{month}'."
        else:
            print(f"Error updating cell: {update_response.status_code} - {update_response.text}")
            return f"Error updating cell: {update_response.status_code} - {update_response.text}"

    except Exception as e:
        print(f"An error occurred: {str(e)}")
        return f"An error occurred: {str(e)}"


def increment_invoice_number(access_token, drive_id, file_path, target_email):
    """
    Increment the "Invoice Number" for a specific email in the "Email" sheet of the Excel file on SharePoint.

    This function dynamically fetches the range of rows containing data in the "Email" sheet to find the target email,
    increments the corresponding invoice number in the "Invoice Number" column, and updates it directly using the
    Microsoft Graph API.

    Args:
        access_token (str): OAuth2 access token for authentication.
        drive_id (str): The ID of the SharePoint drive.
        file_path (str): The relative path to the file in SharePoint.
        target_email (str): The email address to match for incrementing the invoice number.

    Returns:
        str: Status message indicating success or failure.
    """
    try:
        # Get the file ID using the file path
        get_file_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{file_path}"
        headers = {"Authorization": f"Bearer {access_token}"}
        response = requests.get(get_file_url, headers=headers)

        if response.status_code != 200:
            return f"Error fetching file ID: {response.status_code} - {response.text}"

        file_id = response.json().get("id")
        worksheet_name = "Email"  # Correct worksheet name

        # Fetch the used range to dynamically get rows with data
        used_range_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{file_id}/workbook/worksheets('{worksheet_name}')/usedRange"
        used_range_response = requests.get(used_range_url, headers=headers)

        if used_range_response.status_code != 200:
            return f"Error fetching used range: {used_range_response.status_code} - {used_range_response.text}"

        used_range_values = used_range_response.json().get("values", [])
        if not used_range_values:
            return "Error: No data found in the sheet."

        # Find the row containing the target email
        row_index = None
        for index, row in enumerate(used_range_values, start=1):  # 1-based index
            if row and row[0] == target_email:
                row_index = index
                break

        if row_index is None:
            return f"Error: Email '{target_email}' not found in the 'Email' sheet."

        # Increment the invoice number in column D
        cell_address = f"D{row_index}"
        update_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{file_id}/workbook/worksheets('{worksheet_name}')/range(address='{cell_address}')"

        # Get the current value
        current_value_response = requests.get(update_url, headers=headers)
        if current_value_response.status_code != 200:
            return f"Error fetching current invoice number: {current_value_response.status_code} - {current_value_response.text}"

        current_value = current_value_response.json().get("values", [[0]])[0][0]
        new_value = int(current_value) + 1 if current_value else 1  # Increment by 1

        # Update the value
        update_data = {"values": [[new_value]]}
        update_response = requests.patch(update_url, headers=headers, json=update_data)

        if update_response.status_code == 200:
            return f"Successfully incremented invoice number for '{target_email}' to {new_value}."
        else:
            return f"Error updating cell: {update_response.status_code} - {update_response.text}"

    except Exception as e:
        return f"An error occurred: {str(e)}"


def current_academic_year():
    # Determine the current academic year
    current_date = datetime.now()
    current_year = current_date.year
    current_month = current_date.month

    # Calculate the academic year
    if current_month >= 8:  # August to December
        start_year = current_year
        end_year = current_year + 1
    else:  # January to July
        start_year = current_year - 1
        end_year = current_year

    academic_year = f"{start_year}-{str(end_year)[-2:]}"  # E.g., "2024-25"

    return academic_year

def get_or_create_base_folder_path(access_token, drive_id):
    """
    Dynamically determine the base folder path for the academic year and create the folder
    in SharePoint if it doesn't exist.

    Args:
        access_token (str): OAuth2 access token for Microsoft Graph API.
        drive_id (str): The ID of the SharePoint drive.

    Returns:
        str: The SharePoint folder path.
    """
    academic_year = current_academic_year()
    base_folder_path = f"AEB Financial/{academic_year}/Invoices"

    # Check if the folder exists
    check_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{base_folder_path}"
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(check_url, headers=headers)

    if response.status_code == 404:  # Folder does not exist
        # Create the folder
        create_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{base_folder_path}:/children"
        folder_data = {
            "name": academic_year,
            "folder": {},
            "@microsoft.graph.conflictBehavior": "rename"
        }
        create_response = requests.post(create_url, headers=headers, json=folder_data)
        if create_response.status_code == 201:
            print(f"Folder '{base_folder_path}' created successfully.")
        else:
            raise Exception(f"Error creating folder: {create_response.status_code} - {create_response.text}")
    elif response.status_code != 200:
        raise Exception(f"Error checking folder: {response.status_code} - {response.text}")

    # Return the base folder path
    return base_folder_path


def list_files_by_path(access_token, drive_id, folder_path):
    """
    Lists files and folders in a specified path in OneDrive.

    Args:
        access_token (str): OAuth2 access token for authentication.
        drive_id (str): The ID of the drive to query.
        folder_path (str): The path to the folder in the drive.

    Returns:
        list: A list of dictionaries containing file/folder details (name, type, ID).
    """
    headers = {"Authorization": f"Bearer {access_token}"}
    url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{folder_path}:/children"
    response = requests.get(url, headers=headers)
    
    if response.status_code == 200:
        items = response.json().get("value", [])
        result = []
        for item in items:
            result.append({
                "name": item["name"],
                "type": "folder" if "folder" in item else "file",
                "id": item["id"]
            })
        return result
    else:
        print(f"Error fetching files by path: {response.status_code}")
        print(response.json())
        return []


def last():
    st.session_state.clear()
    
# Function to handle adding a row
def add_row():
    st.session_state.table_data.append({"date": "", "time_hours": "", "activity": "", "amount": ""})

# Function to handle removing a row
def remove_row(index):
    if len(st.session_state.table_data) > index:
        st.session_state.table_data.pop(index)
        
def generate_invoice():
    # Prepare data
    data = {
        "inv_date": st.session_state.inv_date,
        "inv_num": st.session_state.inv_num,
        "centre_num": st.session_state.centre_num,
        "hourly_rate": st.session_state.hourly_rate,
        "inv_total": st.session_state.inv_total,
        "ur_name": st.session_state.ur_name,
        "acc_name": st.session_state.acc_name,
        "branch_name": st.session_state.branch_name,
        "sort_code": st.session_state.sort_code,
        "acc_num": st.session_state.acc_num,
        "utr_num": st.session_state.utr
    }
    
    # Generate document
    doc = replace_placeholders(template_file, data, st.session_state.table_data)
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    # Save file locally
    with open(f"Invoice_{st.session_state.safe_name}.docx", "wb") as f:
        f.write(output.getvalue()) 
  
            
def fill_timesheet(template_path, save_path, session_data):
    """
    Copies the template Excel file, fills it with session details, 
    and saves the filled file to the specified path.

    Args:
        template_path (str): Path to the timesheet template file.
        save_path (str): Path to save the filled timesheet.
        session_data (list): List of session dictionaries with keys 'date', 'time', 'topic', and 'duration'.
    """
    # Copy the template to create a new file
    shutil.copy(template_path, save_path)
    
    # Load the copied file to work on
    workbook = load_workbook(save_path)
    sheet = workbook.active

    # Fill columns A to E in the Excel sheet
    for i, session in enumerate(session_data, start=2):  # Start from row 2 assuming row 1 has headers
        session_date = datetime.strptime(session["date"], "%d-%m-%Y")
        day_of_week = calendar.day_name[session_date.weekday()]  # Get the day of the week
        
        sheet[f"A{i}"] = day_of_week  # Day of the week
        sheet[f"B{i}"] = session["date"]  # Date
        sheet[f"C{i}"] = session["time"]  # Time
        sheet[f"D{i}"] = session["duration"]  # Duration in hours
        sheet[f"E{i}"] = session["topic"]  # Description of Activity

    # Save the modified workbook
    workbook.save(save_path)
    

    

def send_email(sender_email, sender_password, receiver_email, subject, body, invoice_attachment, timesheet_attachment=None):
    msg = EmailMessage()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = subject
    msg.set_content(body, subtype='html')

    # Add mandatory invoice attachment
    with open(invoice_attachment, 'rb') as f:
        file_data = f.read()
        file_name = os.path.basename(invoice_attachment)
        msg.add_attachment(file_data, maintype='application', subtype='octet-stream', filename=file_name)

    # Add optional timesheet attachment if provided
    if timesheet_attachment:
        with open(timesheet_attachment, 'rb') as f:
            file_data = f.read()
            file_name = os.path.basename(timesheet_attachment)
            msg.add_attachment(file_data, maintype='application', subtype='octet-stream', filename=file_name)

    # Use the SMTP server for sending the email
    with smtplib.SMTP('smtp.office365.com', 587) as server:
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)
        

# Convert api_events into an HTML table
def generate_events_table(api_events):
    table_rows = ""
    for event in api_events:
        table_rows += f"""
        <tr>
            <td>{event['title']}</td>
            <td>{event['start']}</td>
            <td>{event['end']}</td>
            <td>{event['location']}</td>
        </tr>
        """
    return f"""
    <table border="1" style="border-collapse: collapse; width: 100%; text-align: left;">
        <thead>
            <tr>
                <th>Title</th>
                <th>Start Time</th>
                <th>End Time</th>
                <th>Location</th>
            </tr>
        </thead>
        <tbody>
            {table_rows}
        </tbody>
    </table>
    """
   

# Generate HTML table for validation results
def generate_validation_table(validation_results):
    table_rows = ""
    for result in validation_results:
        session = result["Session"]
        status = result["Status"]
        event_details = ""
        
        # Add event details if matched
        if result["Event"]:
            event = result["Event"]
            event_details = f"""
                <strong>Title:</strong> {event['title']}<br>
                <strong>Start:</strong> {event['start']}<br>
                <strong>End:</strong> {event['end']}<br>
                <strong>Timezone:</strong> {event['timezone']}
            """
        else:
            event_details = "No Event Found"
        
        table_rows += f"""
        <tr>
            <td>{session['date']}</td>
            <td>{session['time']}</td>
            <td>{session['topic']}</td>
            <td>{session['duration']}</td>
            <td>{status}</td>
            <td>{event_details}</td>
        </tr>
        """
    
    return f"""
    <table border="1" style="border-collapse: collapse; width: 100%; text-align: left;">
        <thead>
            <tr>
                <th>Session Date</th>
                <th>Session Time</th>
                <th>Topic</th>
                <th>Duration (hrs)</th>
                <th>Status</th>
                <th>Event Details</th>
            </tr>
        </thead>
        <tbody>
            {table_rows}
        </tbody>
    </table>
    """
    
    
             
# ========================
# Initialize session state
# ========================
if 'step' not in st.session_state: st.session_state.step = 1
if 'recipients' not in st.session_state: st.session_state.recipients = None    
if 'email' not in st.session_state: st.session_state.email = None    
if 'utr' not in st.session_state: st.session_state.utr = None    

if 'inv_date' not in st.session_state: st.session_state.inv_date = date.today().strftime("%d-%m-%Y")
if 'ur_name' not in st.session_state: st.session_state.ur_name = None
if 'inv_num' not in st.session_state: st.session_state.inv_num = None
if 'centre_num' not in st.session_state: st.session_state.centre_num = None
if 'hourly_rate' not in st.session_state: st.session_state.hourly_rate = None
if 'inv_total' not in st.session_state: st.session_state.inv_tot = 0
if 'acc_name' not in st.session_state: st.session_state.acc_name = None
if 'branch_name' not in st.session_state: st.session_state.branch_name = None
if 'sort_code' not in st.session_state: st.session_state.sort_code = "N/A"
if 'acc_num' not in st.session_state: st.session_state.acc_num = None
if 'table_data' not in st.session_state: st.session_state.table_data = [{"date": "", "time_hours": "", "activity": "", "amount": ""}]
if 'inv_total' not in st.session_state: st.session_state.inv_total = 0
if 'user_data' not in st.session_state: st.session_state.user_data = [''] * 12  # Initialize with 10 empty strings
if 'date' not in st.session_state: st.session_state.date = None
if 'time' not in st.session_state: st.session_state.time = None

if 'safe_name' not in st.session_state: st.session_state.safe_name = None

if "session_data" not in st.session_state: st.session_state.session_data = [{"date": "", "time": ""}]

if "timesheet_save_path" not in st.session_state: st.session_state.timesheet_save_path = None
# ========================
# App
# ========================

ACCESS_TOKEN = acquire_access_token() 
template_file = "resources/template_invoice.docx"    

total_steps = 4 # Define the total number of steps
progress = get_progress(st.session_state.step, total_steps) # Calculate the current progress
# Display the progress bar and percentage
st.write(f"Progress: {progress}%")
st.progress(progress)
st.divider()        

global recipients
if st.session_state.step == 1:
    
    # Logo space
    st.markdown(
        """
        <style>
            .logo {
                display: flex;
                justify-content: center;
                align-items: center;
                padding: 20px;
                border-bottom: 2px solid #f0f0f0;
                margin-bottom: 20px;
            }
        </style>
        """, unsafe_allow_html=True
    )
    st.image("resources/logo_removed_bg - enlarged.png", use_column_width=True)
    # Page title
    st.markdown(
        """
        <h2 style="text-align:center; color:#4CAF50;">Prevista Invoice Submission System</h2>
        """, unsafe_allow_html=True
    )
    st.divider()        
        
    # Input placeholders
    # st.write('### Enter your Prevista Email & "Unique Number" and click "Access" button:')
    
    # Email input
    st.session_state.email = st.text_input("Email", placeholder="Enter your email (e.g. format: yourname@prevista.co.uk)")
    # UTR input
    st.session_state.utr = st.text_input("UTR / URN", placeholder="Enter your UTR (e.g. format: UTR123456)")
    
    
    if st.button("Proceed"):
        st.session_state.recipients = fetch_recipients_from_sharepoint(ACCESS_TOKEN, DRIVE_ID)
        # Filter the row where email matches
        st.session_state.user_data = next(
            (row for row in st.session_state.recipients if row[0] == st.session_state.email and row[1] == st.session_state.utr),
            None  # Default value if no match found
        )
        if st.session_state.user_data == None:
            st.error("Email and/or UTR not found. Please check your credentials and try again.")
            st.stop()
        else:
            st.session_state.step=2
            st.experimental_rerun()

elif st.session_state.step == 2:

    # inv_date = st.date_input(
    #             label="Invoice Date",  # Label for the field
    #             value=date.today(),
    #             min_value=date(1900, 1, 1),  # Minimum selectable date
    #             max_value=date.today(),  # Maximum selectable date
    #             help="Choose a date",  # Tooltip text
    #             
    #         )

    st.session_state.inv_date = date.today().strftime("%d-%m-%Y")
    st.write(f"Invoice Date: **{st.session_state.inv_date}**")

    st.session_state.ur_name = st.text_input("Your Name", st.session_state.user_data[2], placeholder="Kasia Kwiatkowska")
    st.session_state.inv_num = st.text_input("Invoice Number", st.session_state.user_data[3], placeholder="74")
    st.session_state.centre_num = st.text_input("Centre Number", st.session_state.user_data[4], placeholder="67890")
    st.session_state.hourly_rate = st.text_input("Pay Rate", st.session_state.user_data[5], placeholder="50")
    # st.session_state.inv_total = st.text_input("Invoice Total", st.session_state.user_data[5], placeholder="500")
    st.session_state.acc_name = st.text_input("Account Holder's Name", st.session_state.user_data[6], placeholder="Kasia Kwiatkowska")
    st.session_state.branch_name = st.text_input("Bank Name", st.session_state.user_data[7], placeholder="ABC Bank")
    st.session_state.sort_code = st.text_input("Sort Code (optional for Non UK)", st.session_state.user_data[8], placeholder="00-00-00")
    st.session_state.acc_num = st.text_input("Account Number / IBAN", st.session_state.user_data[9], placeholder="12345678")

    if st.button("Next"):
        if not all([st.session_state.ur_name, st.session_state.inv_num, st.session_state.centre_num, st.session_state.hourly_rate, st.session_state.acc_name, st.session_state.branch_name, st.session_state.acc_num]):
            st.error("Please fill in all required fields.")
            st.stop()
        else:
            st.session_state.safe_name = re.sub(r'\W+', '_', st.session_state.user_data[2])
            st.session_state.step=3
            st.experimental_rerun()

elif st.session_state.step == 3:

    # Allow users to input table data
    if 'toast_shown' not in st.session_state:
        st.toast('This page is for INVOICE DETAILS ONLY!', icon="🚨")
        st.toast('This page is not for TIMESHEET!', icon="🚨")
        st.session_state.toast_shown = True

    st.write("### Add INVOICE Details")
    for i, row in enumerate(st.session_state.table_data):
        cols = st.columns(4)
        # row["date"] = cols[0].date_input("Date / Period", datetime.strptime(row["date"], "%d-%m-%Y").date() if row["date"] else None, key=f"date_{i}", format='DD/MM/YYYY')
        # if row["date"]!=None:
        #     row["date"] = row["date"].strftime("%d-%m-%Y")
            
        row["date"] = cols[0].text_input("Date / Period", row["date"] if row["date"] else "", key=f"date_{i}", placeholder="")                    
        row["time_hours"] = cols[1].text_input("Time / Hours", row["time_hours"], key=f"time_hours_{i}")
        row["activity"] = cols[2].text_input("Activity / Service Provided", row["activity"], key=f"activity_{i}")
        row["amount"] = cols[3].number_input("Amount (£)", value=float(row["amount"]) if row["amount"] else 0.0, min_value=0.0, key=f"amount_{i}", format="%.2f")        
        if cols[0].button("Remove Row", key=f"remove_row_{i}"):
            remove_row(i)
            st.experimental_rerun()

    # Button to add a new row
    if st.button("Add Row"):
        add_row()
        st.experimental_rerun() 

    # Display table data as a preview
    # st.write("### Work Data Preview")
    # st.write(pd.DataFrame(st.session_state.table_data))

    # Calculate total
    st.session_state.inv_total = sum(float(row["amount"]) for row in st.session_state.table_data)
    st.text(f"Total Amount: £{st.session_state.inv_total:.2f}")
    
    if st.button("Next"):
        st.session_state.step=4
        st.experimental_rerun()
    if st.button("Back"):
        st.session_state.step=2
        st.experimental_rerun()
    

elif st.session_state.step == 4:
        
    receipt_files = st.file_uploader(
        "Upload Expense Receipts (Optional, multiple)", 
        type=["jpg", "png", "pdf"], 
        accept_multiple_files=True,
        help="Upload your receipts here (optional)"
    )
    # Session Details for Tutors
    if st.session_state.user_data[10] == "Tutor":
        st.write("### Add Details for your TIMESHEET")
        st.text("Date & Time should match your calendar for accurate payment.")

        # Single Timezone Selection
        if "timezone" not in st.session_state:
            st.session_state.timezone = "UTC"  # Default timezone

        st.session_state.timezone = st.selectbox(
            "Select Time Zone (applies to all sessions)",
            options=all_timezones,
            index=all_timezones.index(st.session_state.timezone)
        )

        # Dynamic Session Input
        for i, session in enumerate(st.session_state.session_data):
            cols = st.columns(3)  # Adjust to 3 columns for Date, Time, and Duration
            
            # Date Input
            session["date"] = cols[0].date_input(
                "Session Date", 
                datetime.strptime(session["date"], "%d-%m-%Y").date() if session["date"] else None, 
                key=f"session_date_{i}", 
                format='DD/MM/YYYY'
            )
            if session["date"]:
                session["date"] = session["date"].strftime("%d-%m-%Y")
            
            # Time Input
            session["time"] = cols[1].time_input(
                "Session Time (24 Hr Format)", 
                value=datetime.strptime(session["time"], "%H:%M:%S").time() if session["time"] else datetime.min.time(), 
                key=f"session_time_{i}"
            ).strftime("%H:%M:%S")
            
            # Duration Input
            session["duration"] = cols[2].number_input(
                "Duration (hours)", 
                value=float(session.get("duration", 0)), 
                min_value=0.0, step=0.5, 
                key=f"session_duration_{i}"
            )
            
            # Description of Activity (New Line)
            session["topic"] = st.text_input(
                f"Description of Activity (Session {i+1})",  # Dynamic label for clarity
                value=session.get("topic", ""),
                key=f"session_description_{i}"
            )

            # Button to remove session
            # Remove Session Button (New Line)
            if st.button("Remove Session", key=f"remove_session_{i}"):
                st.session_state.session_data.pop(i)
                st.experimental_rerun()

        # Add a new session
        if st.button("Add Session"):
            st.session_state.session_data.append({"date": "", "time": "", "topic": "", "duration": 0})
            st.experimental_rerun()

        # Preview of session data
        # st.write("### Session Data Preview")
        # st.write(pd.DataFrame(st.session_state.session_data))

    # Validation before generating invoice
    if st.button("Submit"):
        with st.spinner("Processing your invoice submission..."):
            
            if st.session_state.user_data[10] == "Tutor":
                # Validate if all session details are filled
                incomplete_sessions = [
                    session for session in st.session_state.session_data 
                    if not session["date"] or not session["time"] or not session["topic"] or session["duration"] <= 0
                ]

                if incomplete_sessions:
                    st.error("Please ensure all session dates and times are filled before generating the invoice.")
                    st.stop()
                else:
                    # Generate timesheet
                    ####################
                    timesheet_template_path = 'resources/template_timesheet.xlsx'  # Replace with your actual template path
                    st.session_state.timesheet_save_path = f'Timesheet_{st.session_state.safe_name}.xlsx'  # Path to save the filled file

                    # st.text("Log: Generatting TimeSheet")
                    fill_timesheet(timesheet_template_path, st.session_state.timesheet_save_path, st.session_state.session_data)
                    
                    # Fetch calendar API events
                    ###########################
                    api_events = fetch_calendar_events(access_token=ACCESS_TOKEN, employee_email=st.session_state.email)
                    # st.write("### Calendar API Events")
                    # st.write(api_events)
                    validation_results = validate_sessions(st.session_state.session_data, api_events, st.session_state.timezone)
                    # st.write("### Validation Results")
                    # st.write(validation_results)            

                
            ##############################
            # All Logic here from sharepoint file upload to update master sheet & Folder creation if doesn't exist
            ##############################

            # Generate Invoice
            # st.text("Log: Generatting Invoice")
            generate_invoice()  # f"Invoice_{safe_name}.docx"
            
            # Get or create the base folder path
            try:
                BASE_FOLDER_PATH = get_or_create_base_folder_path(ACCESS_TOKEN, DRIVE_ID)
                print(f"Base folder path: {BASE_FOLDER_PATH}")                              # "AEB Financial/2024-25/Invoices"
            except Exception as e:
                print(f"    An error occurred: {e}")

            month_folder = get_or_create_month_folder(ACCESS_TOKEN, DRIVE_ID, BASE_FOLDER_PATH) # "Jan-24" or "Feb-24" etc.
            FOLDER_PATH = f"{BASE_FOLDER_PATH}/{month_folder}/Catalyst"

            # Process employee folder
            ##########################
            
            # Option files:
            # * Receipt file 
            # * Timesheet for tutors only
            # Start with receipt files in optional_files
            optional_files = list(receipt_files)

            # Add the timesheet to optional_files if it exists
            if st.session_state.timesheet_save_path and os.path.exists(st.session_state.timesheet_save_path):
                with open(st.session_state.timesheet_save_path, "rb") as file:
                    timesheet_content = BytesIO(file.read())  # Read the file content into BytesIO
                    timesheet_content.name = os.path.basename(st.session_state.timesheet_save_path)  # Retain original file name
                    optional_files.append(timesheet_content)
            
            # Upload to Sharepoint
            process_employee_folder_result_message = process_employee_folder(
                ACCESS_TOKEN,
                DRIVE_ID,
                FOLDER_PATH,
                st.session_state.ur_name,
                f"Invoice_{st.session_state.safe_name}.docx",  # Mandatory invoice
                optional_files=optional_files
            )
            
            # Log for eamil!
            # st.text("Log: Processing employee folder")
            # for message in process_employee_folder_result_message:
            #     st.text(f"Log: (Processing employee folder) {message}")
                
            # Process master sheet
            ######################
            master_FILE_PATH = find_master_sheet_path(ACCESS_TOKEN, DRIVE_ID, f"AEB Financial/{current_academic_year()}")
            EMPLOYEE_NAME = st.session_state.ur_name
            TOTAL = st.session_state.inv_total
            MONTH =  datetime.now().strftime("%b-%y")

            update_mastersheet_message = update_mastersheet_sharepoint(ACCESS_TOKEN, DRIVE_ID, master_FILE_PATH, EMPLOYEE_NAME, TOTAL, MONTH)
            # Log for eamil!
            # st.text("Log: Process master sheet")
            # st.text("Log: "+f"{update_mastersheet_message}")
            
            increment_invoice_number_message = increment_invoice_number(ACCESS_TOKEN, DRIVE_ID, master_FILE_PATH, st.session_state.email)
            # Log for eamil!
            # st.text("Log: Increment invoice number")
            # st.text("Log: "+f"{increment_invoice_number_message}")

            # Email Logs
            ##########################
            subject = f"Invoice of {st.session_state.ur_name} has been submitted."
            
            if st.session_state.user_data[10] == "Tutor":
                # Generate the events table HTML
                events_table_html = generate_events_table(api_events)  
                # Generate the validation results table HTML
                validation_table_html = generate_validation_table(validation_results)              
                
                body = f"""
                <html>
                <body>
                    <p>Invoice Received for <b>Tutor</b> : {st.session_state.ur_name} [{st.session_state.email}],</p>
                    <p>Invoice total: £{st.session_state.inv_total}.</p>
                    <br>
                    
                    <p><b>Log: Move Files to Sharepoint folder: </b></p>
                    {process_employee_folder_result_message}
                    <br>
                    
                    <p><b>Log: Process master sheet: </b></p>
                    {update_mastersheet_message}
                    <br>
                    
                    <p><b>Log: Increment invoice number: </b></p>
                    {increment_invoice_number_message}
                    <br>                    
                    
                    <p><b>Calendar Validation Results:</b></p>
                    {validation_table_html}
                    <br>
                    <p><b>All Fetched Events from Employee Calendar:</b></p>
                    {events_table_html} 

                </body>
                </html>
                """            
                send_email(os.getenv('EMAIL'), os.getenv('PASSWORD'), os.getenv('EMAIL'), subject, body, f"Invoice_{st.session_state.safe_name}.docx", st.session_state.timesheet_save_path)

            else:
                body = f"""
                <html>
                <body>
                    <p>Invoice Received for <b>Employee</b> : {st.session_state.ur_name} [{st.session_state.email}],</p>
                    <p>Invoice total: £{st.session_state.inv_total}.</p>
                    <br>

                    <p><b>Log: Move Files to Sharepoint folder: </b></p>
                    {process_employee_folder_result_message}
                    <br>
                    
                    <p><b>Log: Process master sheet: </b></p>
                    {update_mastersheet_message}
                    <br>
                    
                    <p><b>Log: Increment invoice number: </b></p>
                    {increment_invoice_number_message}
                    <br>                    
                                        
                </body>
                </html>
                """
                send_email(os.getenv('EMAIL'), os.getenv('PASSWORD'), os.getenv('EMAIL'), subject, body, f"Invoice_{st.session_state.safe_name}.docx")
        
        
        # Success & Download Buttons
        ##############################
        
        st.success("Invoice Submitted Successfully!")
        # st.download_button(
        #     label="Download Invoice (Optional)",
        #     data=open(f"Invoice_{st.session_state.safe_name}.docx", "rb"),
        #     file_name=f"Invoice_{st.session_state.safe_name}.docx",
        #     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        # )        
        # if st.session_state.user_data[10] == "Tutor":
        #     # Download button for Timesheet
        #     st.download_button(
        #         label="Download Timesheet (Optional)", 
        #         data=open(f"Timesheet_{st.session_state.safe_name}.xlsx", "rb"),
        #         file_name=f"Timesheet_{st.session_state.safe_name}.xlsx",
        #         mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        #     )         
        
        # Prepare files for download
        if st.session_state.user_data[10] == "Tutor":
            # Create a zip file containing both invoice and timesheet
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zf:
                # Add Invoice
                with open(f"Invoice_{st.session_state.safe_name}.docx", "rb") as invoice_file:
                    zf.writestr(f"Invoice_{st.session_state.safe_name}.docx", invoice_file.read())
                
                # Add Timesheet
                with open(f"Timesheet_{st.session_state.safe_name}.xlsx", "rb") as timesheet_file:
                    zf.writestr(f"Timesheet_{st.session_state.safe_name}.xlsx", timesheet_file.read())
            
            zip_buffer.seek(0)
            st.download_button(
                label="Download Invoice + Timesheet",
                data=zip_buffer,
                file_name=f"Invoice_and_Timesheet_{st.session_state.safe_name}.zip",
                mime="application/zip",
            )
        else:
            # Single file download for invoice only
            st.download_button(
                label="Download Invoice",
                data=open(f"Invoice_{st.session_state.safe_name}.docx", "rb"),
                file_name=f"Invoice_{st.session_state.safe_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        ##############################
        # ####################
        ##############################